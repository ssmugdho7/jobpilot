import os
import re

from pypdf import PdfReader

from app.paths import UPLOAD_DIR


def extract_text_from_pdf(path: str) -> str:
    reader = PdfReader(path)
    parts = []
    for page in reader.pages:
        try:
            parts.append(page.extract_text() or "")
        except Exception:
            continue
    return "\n".join(parts)


def extract_text_from_docx(path: str) -> str:
    from docx import Document

    doc = Document(path)
    parts = []
    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text)
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def extract_text(path: str) -> str:
    ext = os.path.splitext(path)[1].lower()
    if ext == ".pdf":
        return extract_text_from_pdf(path)
    if ext in (".docx", ".doc"):
        return extract_text_from_docx(path)
    raise ValueError(f"Unsupported file type: {ext}")


def save_upload(file_storage, filename: str) -> str:
    dest = os.path.join(UPLOAD_DIR, filename)
    file_storage.save(dest)
    return dest


# -------- lightweight regex fallback (used when no Gemini key) --------

EMAIL_RE = re.compile(r"[\w.+-]+@[\w-]+\.[\w.-]+")
PHONE_RE = re.compile(r"(?:\+?\d[\d\s\-()]{7,}\d)")
URL_RE = re.compile(r"https?://[^\s]+")


def extract_email(text: str) -> str:
    match = EMAIL_RE.search(text or "")
    return match.group(0) if match else ""


def extract_contact(text: str) -> dict:
    emails = EMAIL_RE.findall(text)
    phones = PHONE_RE.findall(text)
    urls = URL_RE.findall(text)

    linkedin = github = portfolio = ""
    for u in urls:
        u_low = u.rstrip(".,;)").lower()
        if "linkedin.com" in u_low and not linkedin:
            linkedin = u.rstrip(".,;)")
        elif "github.com" in u_low and not github:
            github = u.rstrip(".,;)")
        elif not portfolio:
            portfolio = u.rstrip(".,;)")

    return {
        "email": emails[0] if emails else "",
        "phone": phones[0].strip() if phones else "",
        "linkedin": linkedin,
        "github": github,
        "portfolio": portfolio,
    }


def first_line_name(text: str) -> str:
    for line in text.splitlines():
        line = line.strip()
        if line and len(line) < 60 and not re.match(r"^[\W_]+$", line):
            return line
    return ""
