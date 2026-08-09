"""Personalized CV rendering that PRESERVES the uploaded CV's formatting.

The generated CV is a copy of the user's uploaded CV template with only the
text content swapped (summary/skills/experience/education tailored to the job).
Paragraph styles, fonts, colors, spacing and layout stay exactly as uploaded.

- DOCX upload  -> used directly as the template.
- PDF upload   -> converted to DOCX via Word (COM), then used as the template.
- No Word / no upload -> falls back to a clean single-column render.
"""

import hashlib
import os
import subprocess

from app.paths import CV_DIR

# Microsoft Word SaveAs format codes
_WD_DOCX = 16
_WD_PDF = 17

_CONVERT_SCRIPT = os.path.join(os.path.dirname(__file__), "word_convert.ps1")

_WORD_AVAILABLE = None


def _cv_stem(job_id: int) -> str:
    return f"ATS_CV_job_{job_id}"


# ---------------------------------------------------------------------------
# Word availability + conversions
# ---------------------------------------------------------------------------

def word_available() -> bool:
    global _WORD_AVAILABLE
    if _WORD_AVAILABLE is None:
        if os.name != "nt":
            _WORD_AVAILABLE = False
            return False
        try:
            import winreg
            with winreg.OpenKey(winreg.HKEY_LOCAL_MACHINE, r"SOFTWARE\Classes\Word.Application"):
                _WORD_AVAILABLE = True
        except (OSError, ImportError):
            _WORD_AVAILABLE = False
    return _WORD_AVAILABLE


def _run_convert(src: str, dst: str, save_as_code: int, timeout: int = 120) -> bool:
    if not os.path.exists(_CONVERT_SCRIPT):
        return False
    try:
        proc = subprocess.run(
            [
                "powershell", "-NoProfile", "-ExecutionPolicy", "Bypass",
                "-File", _CONVERT_SCRIPT,
                "-In", src,
                "-Out", dst,
                "-Format", str(save_as_code),
            ],
            capture_output=True,
            timeout=timeout,
        )
        return proc.returncode == 0 and os.path.exists(dst) and os.path.getsize(dst) > 0
    except Exception as e:
        print(f"  [render] conversion failed: {e}")
        return False


def _ensure_docx_template(template_path: str | None) -> str | None:
    """Return a DOCX template path derived from the uploaded CV, or None."""
    if not template_path or not os.path.exists(template_path):
        return None
    ext = os.path.splitext(template_path)[1].lower()
    if ext == ".docx":
        return template_path
    if not word_available():
        return None
    digest = hashlib.md5(os.path.abspath(template_path).encode("utf-8")).hexdigest()[:10]
    conv = os.path.join(CV_DIR, f"template_{digest}.docx")
    if os.path.exists(conv) and os.path.getsize(conv) > 0:
        return conv
    if _run_convert(template_path, conv, _WD_DOCX):
        return conv
    return None


def _docx_to_pdf(docx_path: str, pdf_path: str) -> bool:
    if not word_available():
        return False
    return _run_convert(docx_path, pdf_path, _WD_PDF)


# ---------------------------------------------------------------------------
# Text swapping while preserving formatting
# ---------------------------------------------------------------------------

_SECTION_ALIASES = {
    "summary": [
        "professional summary", "summary", "profile", "personal profile",
        "about", "about me", "objective", "career objective", "overview",
    ],
    "skills": [
        "technical skills", "skills", "core skills", "key skills",
        "technologies", "tools", "competencies", "tech stack", "skill",
    ],
    "experience": [
        "work experience", "professional experience", "experience",
        "employment history", "employment", "work history", "career history",
        "internship", "internships",
    ],
    "education": [
        "education", "academic", "academics", "qualification", "qualifications",
        "educational background", "academic background",
    ],
}


def _normalize(text: str) -> str:
    return " ".join((text or "").strip().strip(":#;*•-–").split()).lower()


def _looks_like_heading(norm: str) -> bool:
    if not norm or len(norm) > 50:
        return False
    letters = [c for c in norm if c.isalpha()]
    if not letters:
        return False
    upper = sum(1 for c in letters if c.isupper())
    return upper / len(letters) >= 0.6


def _heading_matches(norm: str, aliases: list[str]) -> bool:
    """Heading detection: exact match, alias-prefix, or contains for
    distinctive long aliases on short/uppercase headings only."""
    if norm in aliases:
        return True
    for alias in aliases:
        if norm.startswith(alias):
            nxt = norm[len(alias):]
            if not nxt or nxt[0] in " :;|":
                return True
        if len(alias) >= 6 and alias in norm and _looks_like_heading(norm):
            return True
    return False


def _segment_sections(doc) -> dict:
    """Return {section: [paragraph, ...]} using heading detection.

    Only TOP-LEVEL body paragraphs are considered, so tables (contact/sidebar
    layouts) are never treated as section content and are never modified.
    """
    sections: dict[str, list] = {}
    current = None
    for para in doc.paragraphs:
        norm = _normalize(para.text)
        if not norm:
            continue
        matched = None
        for section, aliases in _SECTION_ALIASES.items():
            if _heading_matches(norm, aliases):
                matched = section
                break
        if matched:
            current = matched
            sections.setdefault(current, [])
        elif current and len(norm) < 300:
            sections.setdefault(current, []).append(para)
    return sections


def _replace_paragraph_text(para, text: str) -> None:
    """Set paragraph text keeping the first run's formatting (clears the rest)."""
    runs = para.runs
    if not runs:
        para.add_run(text)
        return
    runs[0].text = text
    for extra in runs[1:]:
        extra.text = ""


def _replace_section(sections: dict, section: str, lines: list[str]) -> bool:
    """Swap a section's text line-by-line WITHOUT changing the format.

    Each new line maps onto one existing paragraph of the section, so every
    paragraph keeps its own style (fonts, bullets, numbering, colors, spacing).
    When the new content has more lines, extra paragraphs are created by copying
    the section's first paragraph (so bullets/numbering replicate); when it has
    fewer, the surplus paragraphs are removed. Nothing is restyled.
    """
    from docx.text.paragraph import Paragraph
    import copy

    paras = sections.get(section)
    if not paras:
        return False
    lines = [str(ln).strip() for ln in (lines or []) if str(ln).strip()]
    if not lines:
        return False

    first = paras[0]
    parent = first._parent
    anchor = first._p

    while len(paras) < len(lines):
        clone = copy.deepcopy(first._p)
        anchor.addnext(clone)
        anchor = clone
        paras.append(Paragraph(clone, parent))

    for extra in paras[len(lines):]:
        extra._p.getparent().remove(extra._p)

    for para, line in zip(paras[: len(lines)], lines):
        _replace_paragraph_text(para, line)
    return True


def _as_lines(text: str) -> list[str]:
    out: list[str] = []
    for ln in (text or "").splitlines():
        ln = ln.strip().lstrip("•-–*·")
        if ln:
            out.append(ln)
    return out


def _personalize_docx(template_path: str, out_path: str, profile: dict, ats: dict) -> None:
    from docx import Document

    doc = Document(template_path)
    sections = _segment_sections(doc)

    replacements = {
        "summary": [ats.get("summary") or profile.get("summary") or ""],
        "skills": ats.get("skills") or profile.get("skills") or [],
        "experience": _as_lines(ats.get("experience") or profile.get("experience") or ""),
        "education": _as_lines(ats.get("education") or profile.get("education") or ""),
    }
    for section, lines in replacements.items():
        _replace_section(sections, section, lines)

    doc.save(out_path)


# ---------------------------------------------------------------------------
# Fallback renders (clean single-column, used when no template is available)
# ---------------------------------------------------------------------------

def _render_docx_fallback(job_id: int, path: str, profile: dict, ats: dict) -> str:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.6)
        section.right_margin = Inches(0.6)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10)

    name_p = doc.add_paragraph()
    name_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = name_p.add_run(profile.get("name") or "")
    run.bold = True
    run.font.size = Pt(18)

    contact = ", ".join(filter(None, [
        profile.get("phone") or "",
        profile.get("email") or "",
        profile.get("linkedin") or "",
        profile.get("github") or "",
        profile.get("portfolio") or "",
    ]))
    if contact:
        cp = doc.add_paragraph()
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cp.add_run(contact).font.size = Pt(9)

    def heading(text: str):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(text.upper())
        r.bold = True
        r.font.size = Pt(11)
        r.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

    def body(text: str):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        p.add_run(text)

    heading("Professional Summary")
    body(ats.get("summary") or profile.get("summary") or "")

    heading("Skills")
    skills = ats.get("skills") or profile.get("skills") or []
    if skills:
        doc.add_paragraph(", ".join(skills)).paragraph_format.space_after = Pt(2)

    heading("Experience")
    body(ats.get("experience") or profile.get("experience") or "")

    heading("Education")
    body(ats.get("education") or profile.get("education") or "")

    doc.save(path)
    return path


def _render_pdf_fallback(job_id: int, path: str, profile: dict, ats: dict) -> str:
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.lib.enums import TA_CENTER
    from reportlab.platypus import SimpleDocTemplate, Paragraph
    from reportlab.lib import colors

    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        "CVTitle", parent=styles["Title"], fontSize=18, leading=22,
        alignment=TA_CENTER, spaceAfter=2,
    )
    contact_style = ParagraphStyle(
        "CVContact", parent=styles["Normal"], fontSize=9, leading=12,
        alignment=TA_CENTER, spaceAfter=8,
    )
    heading_style = ParagraphStyle(
        "CVHeading", parent=styles["Heading2"], fontSize=11, leading=13,
        spaceBefore=8, spaceAfter=2, textColor=colors.HexColor("#1F4E79"),
    )
    body_style = ParagraphStyle(
        "CVBody", parent=styles["Normal"], fontSize=10, leading=13, spaceAfter=2,
    )

    def esc(text: str) -> str:
        return (text or "").replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")

    story = []
    story.append(Paragraph(esc(profile.get("name") or ""), title_style))
    contact = ", ".join(filter(None, [
        profile.get("phone") or "",
        profile.get("email") or "",
        profile.get("linkedin") or "",
        profile.get("github") or "",
        profile.get("portfolio") or "",
    ]))
    if contact:
        story.append(Paragraph(esc(contact), contact_style))

    story.append(Paragraph("Professional Summary", heading_style))
    story.append(Paragraph(esc(ats.get("summary") or profile.get("summary") or ""), body_style))

    story.append(Paragraph("Skills", heading_style))
    skills = ats.get("skills") or profile.get("skills") or []
    if skills:
        story.append(Paragraph(esc(", ".join(skills)), body_style))

    story.append(Paragraph("Experience", heading_style))
    story.append(Paragraph(esc(ats.get("experience") or profile.get("experience") or ""), body_style))

    story.append(Paragraph("Education", heading_style))
    story.append(Paragraph(esc(ats.get("education") or profile.get("education") or ""), body_style))

    doc = SimpleDocTemplate(
        path, pagesize=letter,
        topMargin=0.5 * inch, bottomMargin=0.5 * inch,
        leftMargin=0.6 * inch, rightMargin=0.6 * inch,
    )
    doc.build(story)
    return path


# ---------------------------------------------------------------------------
# Public API
# ---------------------------------------------------------------------------

def render_cv_docx(job_id: int, profile: dict, ats: dict, template_path: str | None = None) -> str:
    path = os.path.join(CV_DIR, _cv_stem(job_id) + ".docx")
    tpl = _ensure_docx_template(template_path)
    if tpl:
        try:
            _personalize_docx(tpl, path, profile, ats)
            return path
        except Exception as e:
            print(f"  [render] template personalization failed ({e}); falling back")
    return _render_docx_fallback(job_id, path, profile, ats)


def render_cv_pdf(job_id: int, profile: dict, ats: dict, template_path: str | None = None) -> str:
    path = os.path.join(CV_DIR, _cv_stem(job_id) + ".pdf")
    docx_path = render_cv_docx(job_id, profile, ats, template_path=template_path)
    if _docx_to_pdf(docx_path, path):
        return path
    return _render_pdf_fallback(job_id, path, profile, ats)


# Backwards-compatible aliases used by tests
render_docx = render_cv_docx
render_pdf = render_cv_pdf
