"""Render a tailored CV as DOCX from AI Studio editor state.

Takes the editor's section content + property overrides and produces a DOCX file,
preserving formatting from the base CV template where possible.
"""

import os
import uuid

from app.paths import CV_DIR


def _hex_to_rgb(hex_color: str):
    """Convert hex color string to RGB tuple."""
    hex_color = hex_color.lstrip("#")
    if len(hex_color) == 3:
        hex_color = "".join(c * 2 for c in hex_color)
    if len(hex_color) != 6:
        return None
    try:
        return tuple(int(hex_color[i:i+2], 16) for i in (0, 2, 4))
    except ValueError:
        return None


def _apply_properties(paragraph, props: dict, base_font_size=10):
    """Apply formatting properties to a paragraph."""
    from docx.shared import Pt, RGBColor

    font_size = props.get("font_size")
    if font_size:
        try:
            for run in paragraph.runs:
                run.font.size = Pt(float(font_size))
        except (ValueError, TypeError):
            pass

    color = props.get("color")
    if color:
        rgb = _hex_to_rgb(color)
        if rgb:
            try:
                from docx.shared import RGBColor as RGBColorCls
                for run in paragraph.runs:
                    run.font.color.rgb = RGBColorCls(*rgb)
            except Exception:
                pass

    spacing = props.get("spacing")
    if spacing is not None:
        try:
            from docx.shared import Pt
            paragraph.paragraph_format.space_after = Pt(float(spacing))
        except (ValueError, TypeError):
            pass


def _add_section_to_doc(doc, title: str, content: str, props: dict):
    """Add a section with heading and body content to the document."""
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    # Add heading
    heading_p = doc.add_paragraph()
    heading_p.paragraph_format.space_before = Pt(10)
    heading_p.paragraph_format.space_after = Pt(4)
    run = heading_p.add_run(title.upper())
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

    _apply_properties(heading_p, props)

    # Add body lines
    lines = [ln.strip() for ln in content.splitlines() if ln.strip()]
    if not lines:
        lines = [content.strip()] if content.strip() else []

    for line in lines:
        body_p = doc.add_paragraph()
        body_p.paragraph_format.space_after = Pt(2)
        body_p.add_run(line)
        _apply_properties(body_p, props)


def _personalize_docx_template(template_path: str, out_path: str, sections: dict, properties: dict):
    """Replace sections in a DOCX template while preserving base formatting."""
    from docx import Document

    doc = Document(template_path)

    # Import the section detection from render.py
    from app.cv.render import _segment_sections, _replace_section, _as_lines

    section_map = _segment_sections(doc)

    # Replace standard sections
    for section_name in ("summary", "skills", "experience", "education"):
        content = sections.get(section_name, "")
        props = properties.get(section_name, {})

        if section_name == "skills":
            if isinstance(content, str):
                lines = [s.strip() for s in content.split(",") if s.strip()]
            else:
                lines = [str(s).strip() for s in (content or []) if str(s).strip()]
        elif section_name == "summary":
            lines = [content] if content else []
        else:
            lines = _as_lines(content)

        if lines and section_name in section_map:
            _replace_section(section_map, section_name, lines)

    # Add custom sections at the end
    custom_sections = sections.get("custom_sections") or []
    if custom_sections:
        # Find a good insertion point (after education, before end)
        last_body_para = None
        for para in doc.paragraphs:
            if para.text.strip():
                last_body_para = para

        for cs in custom_sections:
            title = cs.get("title") or ""
            content = cs.get("content") or ""
            if not title or not content:
                continue

            # Add heading
            heading_p = doc.add_paragraph()
            heading_p.paragraph_format.space_before = Pt(10)
            heading_p.paragraph_format.space_after = Pt(4)
            run = heading_p.add_run(title.upper())
            run.bold = True
            run.font.size = Pt(12)

            # Add content lines
            for line in content.splitlines():
                if line.strip():
                    body_p = doc.add_paragraph()
                    body_p.paragraph_format.space_after = Pt(2)
                    body_p.add_run(line.strip())

    doc.save(out_path)


def _render_fallback(sections: dict, properties: dict, profile: dict, out_path: str):
    """Render a clean single-column DOCX when no template is available."""
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    # Set margins
    for section in doc.sections:
        section.top_margin = Pt(36)
        section.bottom_margin = Pt(36)
        section.left_margin = Pt(48)
        section.right_margin = Pt(48)

    # Name
    name_p = doc.add_paragraph()
    name_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = name_p.add_run(profile.get("name") or "")
    run.bold = True
    run.font.size = Pt(18)

    # Contact
    contact = ", ".join(filter(None, [
        profile.get("phone") or "",
        profile.get("email") or "",
        profile.get("linkedin") or "",
        profile.get("github") or "",
    ]))
    if contact:
        cp = doc.add_paragraph()
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cp.add_run(contact).font.size = Pt(9)

    # Sections
    for section_name in ("summary", "skills", "experience", "education"):
        content = sections.get(section_name, "")
        props = properties.get(section_name, {})

        if not content:
            continue

        heading_p = doc.add_paragraph()
        heading_p.paragraph_format.space_before = Pt(10)
        heading_p.paragraph_format.space_after = Pt(4)
        run = heading_p.add_run(section_name.upper())
        run.bold = True
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

        _apply_properties(heading_p, props)

        if section_name == "skills":
            if isinstance(content, str):
                skills_text = content
            else:
                skills_text = ", ".join(str(s) for s in content if s)
            body_p = doc.add_paragraph()
            body_p.paragraph_format.space_after = Pt(2)
            body_p.add_run(skills_text)
            _apply_properties(body_p, props)
        else:
            lines = [ln.strip() for ln in content.splitlines() if ln.strip()]
            for line in lines:
                body_p = doc.add_paragraph()
                body_p.paragraph_format.space_after = Pt(2)
                body_p.add_run(line)
                _apply_properties(body_p, props)

    # Custom sections
    custom_sections = sections.get("custom_sections") or []
    for cs in custom_sections:
        title = cs.get("title") or ""
        content = cs.get("content") or ""
        if not title or not content:
            continue

        heading_p = doc.add_paragraph()
        heading_p.paragraph_format.space_before = Pt(10)
        heading_p.paragraph_format.space_after = Pt(4)
        run = heading_p.add_run(title.upper())
        run.bold = True
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

        for line in content.splitlines():
            if line.strip():
                body_p = doc.add_paragraph()
                body_p.paragraph_format.space_after = Pt(2)
                body_p.add_run(line.strip())

    doc.save(out_path)


def render_studio_docx(
    template_path: str,
    sections: dict,
    properties: dict,
    profile: dict,
    job_title: str = "",
    company: str = "",
) -> str:
    """Render a tailored CV DOCX from AI Studio editor state.

    Args:
        template_path: Path to the base CV template (DOCX).
        sections: Dict with section content (summary, skills, experience, education, custom_sections).
        properties: Dict with per-section property overrides (font_size, color, spacing).
        profile: User profile dict (name, email, etc.).
        job_title: Job title for the filename.
        company: Company name for the filename.

    Returns:
        Path to the generated DOCX file.
    """
    safe_name = "".join(c for c in f"{job_title}_{company}" if c.isalnum() or c in " _-").strip()
    if not safe_name:
        safe_name = "tailored_cv"
    unique_id = uuid.uuid4().hex[:8]
    out_path = os.path.join(CV_DIR, f"studio_{safe_name}_{unique_id}.docx")

    # Try template-based rendering
    if template_path and os.path.exists(template_path):
        ext = os.path.splitext(template_path)[1].lower()
        if ext == ".docx":
            try:
                _personalize_docx_template(template_path, out_path, sections, properties)
                return out_path
            except Exception as e:
                print(f"  [render-studio] template personalization failed ({e}); using fallback")

    # Fallback: clean single-column render
    _render_fallback(sections, properties, profile, out_path)
    return out_path
